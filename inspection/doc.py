#!/usr/bin/env python
# -*- coding:utf-8 -*-

import os
import commands
import json
import sys
import datetime

from docx import *
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION

date1 = datetime.datetime.now().strftime('%Y-%m-%d-%H:%M:%S')
date2 = datetime.datetime.now().strftime('%Y-%m-%d')
scriptdir = os.getcwd()
environment = sys.argv[1]
man = sys.argv[2]

commands.getoutput("ansible gz-%s --private-key /root/.ssh/id_rsa_devops -t %s/%s-log -u devops -b --become-user=root -m script -a '%s/check.sh'"  %(environment, scriptdir, environment, scriptdir))
#======================================生成 word文档（.docx）======================================================================================
document = Document()

##自定义段落样式
style = document.styles['Body Text 2']
font = style.font
font.name = "Times New Roman"
font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
font.bold = True


title = document.add_heading(u"云平台巡检报告")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

line2 = document.add_paragraph(u'日期： ')
run2 = line2.add_run(date2)
#run2.font.underline = True

line3 = line2.add_run(u'\t\t\t\t\t\t 巡检人： ')
line2.add_run(man)
#run3.font.underline = True

document.add_heading(u'物理机',1)

table = document.add_table(rows=1, cols=7,style = 'Table Grid')
table.autofit =False
w = float(2)/2.0
w2 = float(1)/2.0
w3 = float(2.3)/2.0
table.columns[0].width = Inches(w)
table.columns[1].width = Inches(w3)
table.columns[2].width = Inches(w2)
table.columns[3].width = Inches(w3)

hdr_cells = table.rows[0].cells
hdr_cells[0].text =u'主机名'
hdr_cells[1].text =u'IP'
hdr_cells[2].text =u'CPU'
hdr_cells[3].text =u'内存'
hdr_cells[4].text =u'磁盘'
hdr_cells[5].text =u'网卡'
hdr_cells[6].text =u'硬件故障'

document.add_heading(u'网络',1)

table2 = document.add_table(rows=1, cols=2,style = 'Table Grid')
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text =u'公网出口总带宽(近一天最大峰值)'

document.add_heading(u'硬盘健康状态',1)
tabledisk = document.add_table(rows=1, cols=2,style = 'Table Grid')
tabledisk.autofit =False
hdrdisk_cells = tabledisk.rows[0].cells
hdrdisk_cells[0].text =u'主机名'
hdrdisk_cells[1].text =u'硬盘状态'

document.add_heading(u'OpenStack',1)
table3 = document.add_table(rows=1, cols=2,style = 'Table Grid')
openstack = [u'OpenStack服务',u'OpenStack系统用户',u'OpenStack系统租户']
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text =u'Systemd Service'
for j in range(0,len(openstack)):
    row_cells3 = table3.add_row().cells
    row_cells3[0].text = openstack[j]

document.add_heading(u'Ceph',1)
table4 = document.add_table(rows=1, cols=2,style = 'Table Grid')
#table4.columns[1].width = Inches(w3)
ceph = [u'集群容量使用率(阀值：80%)',u'OSD服务状态',u'OSD使用率(阀值：80%)',u'OSD数量',u'Pool数量',u'PG数量']
hdr_cells4 = table4.rows[0].cells
hdr_cells4[0].text =u'集群健康状态'
for k in range(0,len(ceph)):
    row_cells4 = table4.add_row().cells
    row_cells4[0].text = ceph[k]

document.add_heading(u'麒云平台',1)
table5 = document.add_table(rows=1, cols=2,style = 'Table Grid')
hdr_cells5 = table5.rows[0].cells
hdr_cells5[0].text =u'麒云模块健康状态'

document.add_heading(u'其他服务',1)
table6 = document.add_table(rows=1, cols=2,style = 'Table Grid')
other = [u'mariadb',u'rabbitmq',u'mongodb',u'memcache',u'haproxy',u'keepalived']
hdr_cells6 = table6.rows[0].cells
hdr_cells6[0].text =u'reids'
for l in range(0,len(other)):
    row_cells6 = table6.add_row().cells
    row_cells6[0].text = other[l]

document.add_heading(u'管理节点',1)
table7 = document.add_table(rows=1, cols=2,style = 'Table Grid')
manager = [u'域名解析服务',u'zabbix服务']
hdr_cells7 = table7.rows[0].cells
hdr_cells7[0].text =u'容器运行状态'
for o in range(0,len(manager)):
    row_cells7 = table7.add_row().cells
    row_cells7[0].text = manager[o]

#####################插入数据到表格中#####################
datadir = '%s/%s-log' %(scriptdir,environment)
list1 = os.listdir(datadir)
hostlist = []
cpulist = []
disklist = []
diskhealthlist = []
memlist = []
iplist = []
niclist = []
systemservicelist = []
httplist = []
novalist = []
neutronlist = []
cinderlist = []
cephclusterlist = []
openstackuserlist = []
openstacktenantlist = []
cephhealthlist = []
osdvolumelist = []
osdsevicelist = []
osdnumlist = []
poolnumlist = []
pgnumlist = []
qiyunlist = []
redislist = []
mariadblist = []
rabbitmqlist = []
mongodlist = []
memcachedlist = []
haproxylist = []
keepalivedlist = []

for i in range(0,len(list1)):
    path = os.path.join(datadir,list1[i])
    if os.path.isfile(path):
        with open(path) as f:
            info = json.loads(f.read())["stdout_lines"]
            host = str(info[0])
            cpu = info[1]
            mem = info[2]
            disk = info[3]
            ip = info[4]
            nic = info[5]
            hostlist.append(host)
            cpulist.append(cpu)
            memlist.append(mem)
            disklist.append(disk)
            iplist.append(ip)
            niclist.append(nic)

        if "DiskStatus" in info:
           diskservice = info[info.index(u"DiskStatus") + 1:info.index(u"DiskStatusEnd")]
           diskhealthlist.append(host)
           diskhealthlist.extend(diskservice)
           diskhealthlist.append("location")
        if u"Http服务" in info:
           http = info[info.index(u"Http服务") + 1:info.index(u"Http服务") + 3] 
           httplist.append(http)          
        if u"Openstack服务end" in info:
           systemservice = info[info.index(u"Openstack服务") + 1:info.index(u"Openstack服务end")]
           systemservicelist.append(host)
           systemservicelist.extend(systemservice)
        if u"Nova服务" in info:
           nova = info[info.index(u"Nova服务") + 1:info.index(u"Neutron服务")]
           novalist.append(nova)
        if u"Neutron服务" in info:
           neutron = info[info.index(u"Neutron服务") + 1:info.index(u"Cinder服务")]
           neutronlist.append(neutron)
        if u"Cinder服务" in info:
           cinder = info[info.index(u"Cinder服务") + 1:info.index(u"Openstack系统用户")]
           cinderlist.append(cinder)
        if u"Openstack系统用户" in info:
           openstackuser = info[info.index(u"Openstack系统用户") + 1:info.index(u"Openstack系统租户")]
           openstackuserlist.extend(openstackuser)
        if u"Openstack系统租户" in info:
           openstacktenant = info[info.index(u"Openstack系统租户") + 1:info.index(u"Openstack系统租户end")]
           openstacktenantlist.append(openstacktenant)
        if u"Ceph健康状态" in info:
           cephhealth = info[info.index(u"Ceph健康状态") + 1:info.index(u"Ceph集群容量状态")]
           cephhealthlist.append(cephhealth)
        if u"Ceph集群容量状态" in info:
           cephcluster = info[info.index(u"Ceph集群容量状态") + 1:info.index(u"OSD服务状态")]
           cephclusterlist.append(cephcluster)
        if u"OSD服务状态" in info:
           osdsevice = info[info.index(u"OSD服务状态") + 1:info.index(u"OSD容量状态")]
           osdsevicelist.append(osdsevice)
        if u"OSD容量状态" in info:
           osdvolume = info[info.index(u"OSD容量状态") + 1:info.index(u"OSD容量状态end")]
           osdvolumelist.append(osdvolume)
           osdvolumelist.append("\n")
        if u"OSD数量" in info:
           osdnum = info[info.index(u"OSD数量") + 1]
           osdnumlist.append(osdnum)
        if u"Pool数量" in info:
           poolnum = info[info.index(u"Pool数量") + 1]
           poolnumlist.append(poolnum)
        if u"PG数量" in info:
           pgnum = info[info.index(u"PG数量") + 1]
           pgnumlist.append(pgnum)
        if u"麒云模块" in info:
           qiyun = info[info.index(u"麒云模块") + 1:info.index(u"qy_agent") + 2]
           qiyunlist.append(qiyun)
        if u"redis.service" in info:
           redis = info[info.index(u"redis.service") + 1]
           redislist.append(redis)
        if u"mariadb.service" in info:
           mariadb = info[info.index(u"mariadb.service") + 1]
           mariadblist.append(mariadb)
        if u"rabbitmq-server.service" in info:
           rabbitmq = info[info.index(u"rabbitmq-server.service") + 1]
           rabbitmqlist.append(rabbitmq)
        if u"mongod.service" in info:
           mongod = info[info.index(u"mongod.service") + 1]
           mongodlist.append(mongod)
        if u"memcached.service" in info:
           memcached = info[info.index(u"memcached.service") + 1]
           memcachedlist.append(memcached)
        if u"haproxy.service" in info:
           haproxy = info[info.index(u"haproxy.service") + 1]
           haproxylist.append(haproxy)
        if u"keepalived.service" in info:
           keepalived = info[info.index(u"keepalived.service") + 1]
           keepalivedlist.append(keepalived)

print (qiyunlist, redislist, mariadblist, rabbitmqlist, mongodlist, memcachedlist, haproxylist, keepalivedlist)

for i in range(0,len(hostlist)):
    row_cells = table.add_row().cells
    table.cell(i+1,0).text = hostlist[i]
    table.cell(i+1,1).text = iplist[i]
    table.cell(i+1,2).text = cpulist[i]
    table.cell(i+1,3).text = memlist[i]
    table.cell(i+1,4).text = disklist[i]
    table.cell(i+1,5).text = niclist[i]

systemservice = []
if "Error" in systemservicelist:
    [item for item,v in enumerate(systemservicelist) if v=="Error"]
    for i in str(item):
        systemservice.extend(systemservicelist[int(i)-1-1])
        systemservice.append('\n')
        systemservice.extend(systemservicelist[int(i)-1+1])
        systemservice.append('\n')
    systemstatus = systemservice
    table3.cell(0,1).add_paragraph(text=systemstatus, style="Body Text 2")   
else:
    systemstatus = "OK"
    table3.cell(0,1).text = systemstatus

if "Error" in httplist:
    openstackservice = u"Error:控制节点http服务异常"
elif "Error" in novalist:
    novastatus = novalist[0][novalist[0].index("Error") + 1:]	
    openstackservice = u"Nova 服务 Error:\n %s" %(novastatus)
elif "Error" in neutronlist:
    neutronstatus = neutronlist[0][neutronlist[0].index("Error") + 1:]
    openstackservice = u"Neutron 服务 Error:\n %s" %(neutronstatus) 
elif "Error" in cinderlist:
    cinderstatus = cinderlist[0][cinderlist[0].index("Error") + 1:]
    openstackservice = u"Cinder 服务 Error:\n %s" %(cinderstatus)
else:
    openstackservice = "OK"
if "Error" in openstackservice:
    table3.cell(1,1).add_paragraph(text=openstackservice, style="Body Text 2")
else:
    table3.cell(1,1).text = openstackservice

openstackuserstatuslist = []
if "Error" in openstackuserlist:
     for i in range(0,len(openstackuserlist)):
         openstackuserstatuslist.append(openstackuserlist[i])
         openstackuserstatuslist.append('\n')
         openstackuserstatus = openstackuserstatuslist
else:
     openstackuserstatus = "OK"
if "Error" in openstackuserstatus:
    table3.cell(2,1).add_paragraph(text=openstackuserstatus, style="Body Text 2")
else:
    table3.cell(2,1).text = openstackuserstatus

openstacktenantstatuslist = []
if "Error" in openstacktenantlist:
    for i in range(0,len(openstacktenantlist)):
         openstacktenantstatuslist.append(openstackuserlist[i])
         openstacktenantstatuslist.append('\n')
         openstacktenantstatus = openstacktenantstatuslist
else:
     openstacktenantstatus = "OK"
if "Error" in openstacktenantstatus:
    table3.cell(3,1).add_paragraph(text=openstacktenantstatus, style="Body Text 2")
else:
    table3.cell(3,1).text = openstacktenantstatus


cephhealthstatuslist = []
if "Error" in cephhealthlist[0]:
	for i in range(0,len(cephhealthlist[0])):
         cephhealthstatuslist.append(cephhealthlist[0][i])
         cephhealthstatuslist.append('\n')
         cephclusterhealth = cephhealthstatuslist
else:
    cephclusterhealth = "OK"
table4.cell(0,1).text = cephclusterhealth

if "Error" in cephclusterlist[0]:
     cephclusterstatus = cephclusterlist[0][cephclusterlist[0].index("Error"):]
     cephclustervolume = cephclusterstatus
else:
     cephclustervolume = cephclusterlist[0]
if "Error" in cephclustervolume:
     table4.cell(1,1).add_paragraph(text=cephclustervolume, style="Body Text 2")
else:
     table4.cell(1,1).text = cephclustervolume

if "Error" in osdsevicelist[0]:
     osdsevicestatus = osdsevicelist[0][osdsevicelist[0].index("Error"):]
     osdsystemsevice = osdsevicestatus
else:
     osdsystemsevice = "OK"
table4.cell(2,1).text = osdsystemsevice

if "Error" in osdvolumelist[0]:
     osdvolumestatus = osdvolumelist[0][osdvolumelist[0].index("Error"):]
     osdvolumesevice = osdvolumestatus
else:
     osdvolumesevice = osdvolumelist[0]
if "Error" in osdvolumesevice:
     table4.cell(3,1).add_paragraph(text=osdvolumesevice, style="Body Text 2")
else:
     table4.cell(3,1).text = osdvolumesevice
table4.cell(4,1).text = osdnumlist[0]
table4.cell(5,1).text = poolnumlist[0]
table4.cell(6,1).text = pgnumlist[0]

if "Error" in qiyunlist[0]:
     qiyunstatus = qiyunlist[0][qiyunlist[0].index("Error") + 1:]
     qiyunsevice = u"Error\n %s" %qiyunstatus
     table5.cell(0,1).add_paragraph(text=qiyunsevice, style="Body Text 2")
else:
     qiyunsevice = "OK"
     table5.cell(0,1).text = qiyunsevice

if "Error" in redislist:
     redisstatus = redislist[0]
     table6.cell(0,1).add_paragraph(text=redisstatus, style="Body Text 2")
else:
     redisstatus = "OK"
     table6.cell(0,1).text = redisstatus

if "Error" in mariadblist:
    mariadbstatus = mariadblist[0]
    table6.cell(1,1).add_paragraph(text=mariadbstatus, style="Body Text 2")
else:
    mariadbstatus = "OK"
    table6.cell(1,1).text = mariadbstatus

if "Error" in rabbitmqlist:
    rabbitmqstatus = rabbitmqlist[0]
    table6.cell(2,1).add_paragraph(text=rabbitmqstatus, style="Body Text 2")
else:
    rabbitmqstatus = "OK"
    table6.cell(2,1).text = rabbitmqstatus

if "Error" in mongodlist:
    mongodstatus = mongodlist[0]
    table6.cell(3,1).add_paragraph(text=mongodstatus, style="Body Text 2")
else:
    mongodstatus = "OK"
    table6.cell(3,1).text = mongodstatus

if "Error" in memcachedlist:
    memcachedstatus = memcachedlist[0]
    table6.cell(4,1).add_paragraph(text=memcachedstatus, style="Body Text 2")
else:
    memcachedstatus = "OK"
    table6.cell(4,1).text = memcachedstatus

if len(haproxylist):
    if "Error" in haproxylist:
        haproxystatus = haproxylist
        table6.cell(5,1).add_paragraph(text=haproxystatus, style="Body Text 2")
    else:
        haproxystatus = "OK"
        table6.cell(5,1).text = haproxystatus
else:
     haproxystatus = u'无 haproxy 服务'
     table6.cell(5,1).add_paragraph(text=haproxystatus, style="Body Text 2")

if len(keepalivedlist):
    if "Error" in keepalivedlist:
        keepalivedstatus = keepalivedlist
        table6.cell(6,1).add_paragraph(text=keepalivedstatus, style="Body Text 2")
    else:
        keepalivedstatus = "OK"
        table6.cell(6,1).text = keepalivedstatus
else:
	keepalivedstatus = u'无 keepalived 服务'
table6.cell(6,1).add_paragraph(text=keepalivedstatus, style="Body Text 2")


tmp_list = []
count = 1
new_dic = {}

for i in diskhealthlist:
  if i == "location":
    new_dic[count] = tmp_list
    count = count +1
    tmp_list = []
  else:
    tmp_list.append(i)

w4 = float(7.5)/2.0
tabledisk.columns[1].width = Inches(w4)

for i in range(1,len(new_dic)+1):
    row_cells = tabledisk.add_row().cells
    tabledisk.cell(i,0).text = new_dic[i][0]
    tabledisk.cell(i,1).text = new_dic[i][1:]

status,result = commands.getstatusoutput('python ./net.py')
if status != 0:
     netstatus = u"获取公网出口总带宽失败"
else:
     netstatus = json.loads(result)['peak_value_max']	
table2.cell(0,1).text = u'%s Mbit/s' %(round((float(netstatus)/1024)/1024,2))

managestatus,manageresult = commands.getstatusoutput('bash check.sh')
if managestatus !=0:
     print '管理节点运行巡检脚本失败'
     sys.exit
else:
     managerlist = manageresult.split('\n')

row_cells = table.add_row().cells
table.cell(-1,0).text = unicode(managerlist[0],"utf-8")
table.cell(-1,1).text = unicode(managerlist[4],"utf-8")
table.cell(-1,2).text = unicode(managerlist[1],"utf-8")
table.cell(-1,3).text = unicode(managerlist[2],"utf-8")
table.cell(-1,4).text = unicode(managerlist[3],"utf-8")
table.cell(-1,5).text = unicode(managerlist[5],"utf-8")

managerdisklist = []

for i in managerlist[managerlist.index("DiskStatus") + 1:managerlist.index("DiskStatusEnd")]:
     managerdisklist.append(i)
     managerdisklist.append('\n')

managerdiskstatus = ''.join(managerdisklist)


tabledisk.add_row().cells
tabledisk.cell(-1,0).text = unicode(managerlist[0],"utf-8")
tabledisk.cell(-1,1).text = unicode(managerdiskstatus,"utf-8")

containerlist = []
for i in managerlist[managerlist.index('Container') + 1:managerlist.index('DnsNameServer')]:
     containerlist.append(i)
     containerlist.append('\n')

if "Error" in containerlist:
    table7.cell(0,1).add_paragraph(text=containerlist, style="Body Text 2")
else:
    table7.cell(0,1).text = containerlist

if "Error" in managerlist[managerlist.index('DnsNameServer') + 1]:
    table7.cell(1,1).add_paragraph(text=managerlist[managerlist.index('DnsNameServer') + 1], style="Body Text 2")
else:
    table7.cell(1,1).text = managerlist[managerlist.index('DnsNameServer') + 1]

if "Error" in managerlist[managerlist.index('ZabbixServer') + 1]:
    table7.cell(2,1).add_paragraph(text=managerlist[managerlist.index('ZabbixServer') + 1], style="Body Text 2")
else:
    table7.cell(2,1).text = managerlist[managerlist.index('ZabbixServer') + 1]

document.save(u'%s-%s.docx' %(environment, date1))

##############################   发送邮件  #########################################
content = """
大家好：
   附件为 %s云环境巡检报告，请查收知悉
""" %(environment)

with open('./content.txt','wb') as f:
     f.write(content)    

mailstatus, mailresult = commands.getstatusoutput("python ./sendmail.py larf@vipstack.net,railon.wu@vipstack.net,wei.yang@vipstack.net,eden.long@vipstack.net,alex.lin@vipstack.net,zewei.huang@vipstack.net,sam.wu@vipstack.net,wenxing.wang@vipstack.net  %s云环境巡检报告-%s %s/content.txt %s/%s-%s.docx not" %(environment, date1, scriptdir, scriptdir, environment, date1))
#mailstatus, mailresult = commands.getstatusoutput("python ./sendmail.py zewei.huang@vipstack.net  %s云环境巡检报告-%s %s/content.txt %s/%s-%s.docx not" %(environment, date1, scriptdir, scriptdir, environment, date1))
